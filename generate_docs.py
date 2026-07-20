from docx import Document
identity='David Sam Limbong | III RPLK | 2322101894'
cases=[]
login=['Valid credentials','Unknown username','Wrong password','Empty username','Empty password','Both empty','SQL injection username','SQL injection password','Whitespace username','Case-sensitive username','Case-sensitive password','Very long username','HTML username','Refresh after login','Active session redirect','Session fixation','Repeated submit']
register=['Valid fields','Duplicate username','Mismatched passwords','Empty name','Empty email','Empty username','Empty password','Empty re-password','All empty','Email without @','Email without domain','Long email','Short password','Long password','Long username','Unicode name','SQL injection username','SQL injection email','XSS name','XSS username','Apostrophe name','Whitespace username','Active session','Verify DB name','Password hashed','Duplicate email','Refresh register','Double submit','POST content type','CSRF absent']
for i,x in enumerate(login,1): cases.append((f'LG-{i:02}','Login',x,'Expected behavior'))
for i,x in enumerate(register,1): cases.append((f'RG-{i:02}','Register',x,'Expected behavior'))
d=Document(); d.add_heading('LAPORAN PENGUJIAN MODUL LOGIN DAN REGISTER',0); d.add_paragraph(identity); d.add_paragraph('Aplikasi quiz-pengupil | PHP + MySQL + Selenium + pytest')
d.add_heading('1. Pendahuluan',1); d.add_paragraph('Pengujian meliputi fungsionalitas, validasi, keamanan dasar, session, dan integritas database pada login.php dan register.php.')
d.add_heading('2. Metodologi',1); d.add_paragraph('Selenium WebDriver mengotomasi browser. Fixture stub mereset users dan menanam akun bcrypt sebelum setiap test. Assertion memeriksa URL, pesan UI, session, dan database melalui mysql-connector-python. CI memakai MySQL service, PHP server, Chrome headless, pytest, dan artifact.')
d.add_heading(f'3. Daftar Testcase ({len(cases)} skenario)',1); t=d.add_table(rows=1,cols=4); t.style='Table Grid'
for c,v in zip(t.rows[0].cells,['ID','Modul','Skenario','Expected output']): c.text=v
for row in cases:
    cells=t.add_row().cells
    for c,v in zip(cells,row): c.text=v
d.add_heading('4. Temuan Bug',1); d.add_paragraph('$nama digunakan alih-alih $name; error password salah tidak ditampilkan; session guard memakai key berbeda; validasi email/panjang password belum ada; output belum di-escape (risiko XSS).')
d.add_heading('5. Artefak dan CI/CD',1); d.add_paragraph('tests/conftest.py, db_stub.py, test_login.py, test_register.py, requirements.txt, dan .github/workflows/selenium-ci.yml. Perintah: pytest tests/ -v --junitxml=artifacts/report.xml')
d.add_heading('6. Kesimpulan',1); d.add_paragraph(f'{len(cases)} testcase mencakup skenario positif, negatif, boundary, injection, XSS, session, dan database.')
d.add_heading('7. Repository',1); d.add_paragraph('https://github.com/davidlimss/quiz-pengupil-testing'); d.save('Laporan_Pengujian_Quiz_Pengupil.docx')
print('created docx')
